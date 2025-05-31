import streamlit as st
import os
import re
from PyPDF2 import PdfReader
import docx
from langchain.chat_models import ChatOpenAI
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import FAISS
from langchain.chains import ConversationalRetrievalChain
from langchain.memory import ConversationBufferMemory
from streamlit_chat import message
from langchain.prompts import PromptTemplate
from langchain.schema import Document

PROMPT = None

def main():
    st.set_page_config(page_title="Chat with Organic Docs", page_icon="🌿", layout="wide")

    file_mapping = {
        "land preparation": "land preparation.pdf",
        "Online Store (pakorganic.com)": "web"
    }

    manual_source_mapping = {
        "Online Store (pakorganic.com)": [
            "https://pakorganic.com/",
            "https://pakorganic.com/page/2/",
            "https://pakorganic.com/organic-farming-consultancy/",
            "https://pakorganic.com/our-projects/"
        ]
    }

    file_options = list(file_mapping.keys())

    query_params = st.experimental_get_query_params()
    param_selected = query_params.get("option", [file_options[0]])[0]
    trigger = query_params.get("run", ["false"])[0].lower() == "true"

    selected_label = st.session_state.get("last_processed_file", param_selected)

    st.markdown(f"""
        <style>
            .fixed-header {{
                position: fixed;
                top: 4px;
                left: 0;
                right: 0;
                width: 95%;
                margin: auto;
                background-color: #4CAF50;
                padding: 25px;
                z-index: 999;
                border-radius: 10px 10px 0 0;
            }}
            .dropdown-row {{
                margin-top: 10px;
                display: flex;
                justify-content: center;
                align-items: center;
                gap: 10px;
            }}
            .stApp {{
                margin-top: 160px;
            }}
        </style>
        <div class="fixed-header">
            <h2 style="color: white; text-align:center;">🌿🌱 AI-Powered Chatbot 🤖 for Transition to Organic Farming 🌿🌱</h2>
            <p style="color: white; text-align:center;">Get instant answers from organic farming in the context of Pakistan.</p>
            <div class="dropdown-row">
                <form action="" method="get">
                    <label style="color:white;">Select your area of interest:</label>
                    <select name="option">
                        {''.join([f'<option value="{opt}" {"selected" if selected_label == opt else ""}>{opt}</option>' for opt in file_options])}
                    </select>
                    <input type="hidden" name="run" value="true">
                    <button type="submit">Process</button>
                </form>
            </div>
        </div>
    """, unsafe_allow_html=True)

    st.markdown("""
    <div style='text-align: center; margin-top: 5px; font-size: 18px; color: #333;'>
        👉 <strong>Please select your area of interest from the dropdown above and click "Process" to begin.</strong>
    </div>
    """, unsafe_allow_html=True)

    if "chat_history_messages" not in st.session_state:
        st.session_state.chat_history_messages = []

    if "conversation" not in st.session_state:
        st.session_state.conversation = None
        st.session_state.processComplete = False

    if trigger and selected_label != st.session_state.get("last_processed_file"):
        with st.spinner("Processing..."):
            openai_key = st.secrets["OPENAI_API_KEY"]
            file_key = file_mapping[selected_label]

            if file_key == "web":
                urls = manual_source_mapping[selected_label]
                documents = get_web_documents(urls)

                source_references = "\n".join(urls)
                template = f"""
You are a helpful assistant. Use the following web content as your primary reference to answer the user's question.

If the answer is not clearly available, make a reasonable guess or summarize based on what is available.

If you still cannot find anything useful, say:
'Sorry, this question is out of my knowledge domain.'

Try to respond clearly, especially for questions like:
- What are the available products?
- Show product names and prices.
- Give me a list of items sold.

Always add:
Source: {source_references}

Context:
{{context}}

Question:
{{question}}
"""
            else:
                file_path = os.path.join(os.getcwd(), file_key)
                text = get_file_text(file_path)
                documents = [Document(page_content=text)]

                template = """
You are a helpful assistant. Use ONLY the following context to answer the question.
If the answer is not contained in the context, say 'Sorry, this question is out of my knowledge domain.' Don't include source line in that case.

Context:
{context}

Question:
{question}
"""

            prompt = PromptTemplate(input_variables=["context", "question"], template=template)
            text_chunks = get_text_chunks_with_metadata(documents)
            vectorstore = get_vectorstore_with_metadata(text_chunks)
            st.session_state.conversation = get_conversation_chain(vectorstore, openai_key, prompt)
            st.session_state.processComplete = True
            st.session_state.last_processed_file = selected_label
            st.session_state.chat_history_messages = []
            st.success("You may now chat with the selected content.")

    if st.session_state.processComplete and st.session_state.conversation:
        user_question = st.chat_input("Ask a question about the selected content:")
        if user_question:
            lower_question = user_question.strip().lower()

            if is_small_talk(lower_question):
                answer = get_small_talk_reply(lower_question)
            else:
                with st.spinner("Getting your answer..."):
                    response = st.session_state.conversation({'question': user_question})
                    answer = response.get('answer', "Sorry, I couldn't find an answer to that question.")

            st.session_state.chat_history_messages.append({"role": "user", "content": user_question})
            st.session_state.chat_history_messages.append({"role": "bot", "content": answer})

        for i, chat in enumerate(st.session_state.chat_history_messages):
            message(chat["content"], is_user=(chat["role"] == "user"), key=f"chat_{i}")

# ================= Helper Functions =================

def get_file_text(file_path):
    text = ""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        text += get_pdf_text(file_path)
    elif ext == ".docx":
        text += get_docx_text(file_path)
    return text

def get_pdf_text(file_path):
    with open(file_path, "rb") as f:
        reader = PdfReader(f)
        return "".join([page.extract_text() for page in reader.pages if page.extract_text()])

def get_docx_text(file_path):
    doc = docx.Document(file_path)
    return " ".join([para.text for para in doc.paragraphs])

def get_web_documents(urls):
    from langchain.document_loaders import WebBaseLoader
    loader = WebBaseLoader(urls)
    return loader.load()

def get_text_chunks_with_metadata(documents):
    splitter = RecursiveCharacterTextSplitter(chunk_size=1500, chunk_overlap=200)
    return splitter.split_documents(documents)

def get_vectorstore_with_metadata(documents):
    embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
    return FAISS.from_documents(documents, embeddings)

def get_conversation_chain(vectorstore, openai_api_key, prompt):
    llm = ChatOpenAI(openai_api_key=openai_api_key, model_name='gpt-3.5-turbo', temperature=0)
    memory = ConversationBufferMemory(memory_key='chat_history', return_messages=True)
    return ConversationalRetrievalChain.from_llm(
        llm=llm,
        retriever=vectorstore.as_retriever(search_kwargs={"k": 4}),
        memory=memory,
        combine_docs_chain_kwargs={"prompt": prompt}
    )

def is_small_talk(message):
    small_talk_patterns = [
        r"^thanks?$", r"^thank\s?you$", r"^thankyou$", r"^hello$", r"^helo$",
        r"^thx$", r"^tnx$", r"^thanx$", r"^ok$", r"^okay$", r"^great$",
        r"^awesome$", r"^nice$", r"^cool$", r"^that('?s)? (great|good)$", r"^that was helpful$"
    ]
    return any(re.match(pattern, message) for pattern in small_talk_patterns)

def get_small_talk_reply(message):
    reply_map = {
        "thank you": "You're welcome!",
        "thankyou": "You're welcome!",
        "thanks": "You're welcome!",
        "thx": "You're welcome!",
        "tnx": "You're welcome!",
        "thanx": "You're welcome!",
        "hello": "You're welcome! how can I help you!",
        "helo": "You're welcome! how can I help you!",
        "ok": "👍",
        "okay": "👍",
        "great": "Glad to hear that!",
        "awesome": "Glad to hear that!",
        "nice": "Glad to hear that!",
        "cool": "Glad to hear that!",
        "that is great": "Awesome!",
        "that's great": "Awesome!",
        "that's good": "👍",
        "that was helpful": "I'm glad I could help!"
    }

    for key in reply_map:
        if re.match(rf"^{key}$", message):
            return reply_map[key]
    return "You're welcome!"

if __name__ == '__main__':
    main()
